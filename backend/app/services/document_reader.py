"""
Extractor multi-formato para adjuntos de emails entrantes.

Sprint FF 2026-05-27 — F1: lectura de texto extraíble (sin OCR / visión).

Función pública: `extract_text(content_bytes, filename, content_type) -> str`.
Devuelve el texto del adjunto (truncado a MAX_CHARS_POR_ADJUNTO) o "" si
no se puede leer (formato no soportado / PDF escaneado / error).

Soportados en F1:
- PDF con texto (pdfplumber)
- DOCX (python-docx)
- XLSX (openpyxl)
- TXT / CSV / HTML / Markdown / JSON (decode UTF-8 con fallback latin-1)

NO soportados aún (F2 — Claude visión):
- PDF escaneado (sin texto extraíble)
- Imágenes JPG / PNG / TIFF / etc.

Filosofía: best-effort + log warn ante cualquier falla. NUNCA debe
romper el flow de generación de borrador.
"""

from __future__ import annotations

import base64
import io
import logging
import os
from typing import Iterable

logger = logging.getLogger(__name__)

# Límites duros para evitar prompts gigantes
MAX_CHARS_POR_ADJUNTO = 3000
MAX_CHARS_TOTAL = 8000
MAX_ADJUNTOS = 5

# F2 — Claude visión para PDFs escaneados + imágenes
# Sprint FF 2026-05-27
VISION_MAX_BYTES = 5 * 1024 * 1024            # 5 MB hard limit Anthropic
VISION_MIN_TEXT_CHARS_PDF = 100               # PDF text < N → asumir escaneado
VISION_MODEL = "claude-haiku-4-5-20251001"   # más barato, suficiente para extract
VISION_MAX_TOKENS = 1500
_VISION_IMAGE_TYPES = ("image/jpeg", "image/png", "image/gif", "image/webp")


# ──────────────────────────────────────────────────────────────────────────
# Extractores por formato
# ──────────────────────────────────────────────────────────────────────────

def _extract_pdf(content: bytes) -> str:
    """PDF con texto. PDFs escaneados → string vacío (F2 con visión)."""
    try:
        import pdfplumber
        out: list[str] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages[:20]:  # cap 20 páginas
                txt = page.extract_text() or ""
                if txt.strip():
                    out.append(txt)
        result = "\n\n".join(out).strip()
        return result
    except Exception as e:
        logger.warning("PDF extract falló: %s", e)
        return ""


def _extract_docx(content: bytes) -> str:
    """Microsoft Word (.docx)."""
    try:
        import docx as _docx
        d = _docx.Document(io.BytesIO(content))
        partes: list[str] = []
        # Párrafos
        for p in d.paragraphs:
            if p.text.strip():
                partes.append(p.text)
        # Tablas (case típico: tutela con datos en tabla)
        for tabla in d.tables:
            for row in tabla.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    partes.append(" | ".join(cells))
        return "\n".join(partes).strip()
    except Exception as e:
        logger.warning("DOCX extract falló: %s", e)
        return ""


def _extract_xlsx(content: bytes) -> str:
    """Excel (.xlsx). Solo primera hoja, primeras 200 filas."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True, read_only=True)
        ws = wb.active
        partes: list[str] = [f"Hoja: {ws.title}"]
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i >= 200:
                break
            cells = [str(v) for v in row if v is not None]
            if cells:
                partes.append(" | ".join(cells))
        return "\n".join(partes).strip()
    except Exception as e:
        logger.warning("XLSX extract falló: %s", e)
        return ""


def _extract_text_plain(content: bytes) -> str:
    """TXT / CSV / HTML / MD / JSON — best-effort decode."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return content.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    logger.warning("text plain: no decodificable")
    return ""


# ──────────────────────────────────────────────────────────────────────────
# API pública
# ──────────────────────────────────────────────────────────────────────────

def _detectar_formato(filename: str, content_type: str) -> str:
    """Devuelve 'pdf' | 'docx' | 'xlsx' | 'text' | 'image' | 'unsupported'."""
    fn = (filename or "").lower()
    ct = (content_type or "").lower()

    if fn.endswith(".pdf") or "pdf" in ct:
        return "pdf"
    if fn.endswith(".docx") or "wordprocessingml" in ct:
        return "docx"
    if fn.endswith(".xlsx") or "spreadsheetml" in ct:
        return "xlsx"
    if any(fn.endswith(ext) for ext in (".txt", ".csv", ".html", ".htm",
                                          ".md", ".json", ".log", ".xml")):
        return "text"
    if ct.startswith("text/"):
        return "text"
    # F2: solo formatos de imagen soportados por Anthropic visión
    if any(fn.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".gif", ".webp")):
        return "image"
    if ct in _VISION_IMAGE_TYPES:
        return "image"
    # .doc legacy + TIFF + otros → no soportado
    return "unsupported"


def _image_media_type(filename: str, content_type: str) -> str:
    """Normaliza el media_type para Anthropic vision (image/jpeg|png|gif|webp)."""
    ct = (content_type or "").lower().strip()
    if ct in _VISION_IMAGE_TYPES:
        return ct
    fn = (filename or "").lower()
    if fn.endswith((".jpg", ".jpeg")):
        return "image/jpeg"
    if fn.endswith(".png"):
        return "image/png"
    if fn.endswith(".gif"):
        return "image/gif"
    if fn.endswith(".webp"):
        return "image/webp"
    return "image/jpeg"  # fallback


def _extract_with_vision(content: bytes, filename: str, content_type: str,
                         es_pdf: bool) -> str:
    """Extrae texto vía Claude visión (PDFs escaneados + imágenes).

    Sprint FF F2 2026-05-27 — best-effort. Si falla (API key faltante,
    archivo demasiado grande, Anthropic error), devuelve "" y el caller
    sigue sin texto extraído.

    Costo aproximado por llamada: $0.005-0.05 (Haiku, hasta 1500 tokens out).
    Cap por adjunto: 5MB.
    """
    api_key = os.environ.get("ANTHROPIC_API_KEY")
    if not api_key:
        logger.info("vision: ANTHROPIC_API_KEY ausente — skip")
        return ""

    if len(content) > VISION_MAX_BYTES:
        logger.warning("vision: doc %s muy grande (%d bytes > %d) — skip",
                       filename, len(content), VISION_MAX_BYTES)
        return ""

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)
        b64 = base64.standard_b64encode(content).decode("utf-8")

        if es_pdf:
            doc_block = {
                "type": "document",
                "source": {"type": "base64",
                           "media_type": "application/pdf",
                           "data": b64},
            }
        else:
            mt = _image_media_type(filename, content_type)
            doc_block = {
                "type": "image",
                "source": {"type": "base64", "media_type": mt, "data": b64},
            }

        prompt = (
            "Extrae el texto del siguiente documento (legal u operativo) en "
            "español. Si es una tutela / oficio / decreto / sentencia / auto "
            "judicial, identifica: número de proceso, juzgado, partes (accionante "
            "y accionado), causa o materia, plazos y decisión si la hay. Si es "
            "un comprobante de pago: fecha, monto, beneficiario, referencia. "
            "Si es un documento de identidad: tipo, número, nombre. "
            "Devuelve SOLO el texto/datos extraídos en formato legible, sin "
            "comentarios meta tipo 'aquí está el texto'."
        )

        resp = client.messages.create(
            model=VISION_MODEL,
            max_tokens=VISION_MAX_TOKENS,
            messages=[{
                "role": "user",
                "content": [doc_block, {"type": "text", "text": prompt}],
            }],
        )
        txt = (resp.content[0].text or "").strip() if resp.content else ""
        logger.info("vision OK %s: %d chars extraídos", filename, len(txt))
        return txt
    except Exception as e:
        logger.warning("vision extract falló para %s: %s", filename, e)
        return ""


def extract_text(
    content: bytes,
    filename: str = "",
    content_type: str = "",
    max_chars: int = MAX_CHARS_POR_ADJUNTO,
    *,
    enable_vision: bool = True,
) -> str:
    """Extrae texto de un adjunto. Devuelve "" si no se puede leer.

    Pipeline:
    1. Detecta formato por extension/content_type.
    2. Para PDF/DOCX/XLSX/TEXT: extrae localmente (gratis, rápido).
    3. F2: Si el PDF text extraction devuelve muy poco texto (asumir
       escaneado) Y `enable_vision=True` → Claude visión.
    4. F2: Imágenes (JPG/PNG/GIF/WEBP) → Claude visión directo.

    NUNCA levanta excepciones — best-effort con log warn ante errores.
    """
    if not content:
        return ""

    fmt = _detectar_formato(filename, content_type)
    text = ""

    if fmt == "pdf":
        text = _extract_pdf(content)
        # F2: PDF con poco texto extraíble → probablemente escaneado → visión
        if enable_vision and len(text) < VISION_MIN_TEXT_CHARS_PDF:
            logger.info("PDF %s text-extract dió %d chars — usando visión",
                        filename, len(text))
            vision_text = _extract_with_vision(content, filename, content_type, es_pdf=True)
            if vision_text:
                text = vision_text
    elif fmt == "docx":
        text = _extract_docx(content)
    elif fmt == "xlsx":
        text = _extract_xlsx(content)
    elif fmt == "text":
        text = _extract_text_plain(content)
    elif fmt == "image":
        if enable_vision:
            text = _extract_with_vision(content, filename, content_type, es_pdf=False)
        else:
            logger.info("image %s: vision disabled — skip", filename)
            return ""
    else:
        logger.info("formato no soportado: %s / %s", filename, content_type)
        return ""

    # Truncar y limpiar
    text = (text or "").strip()
    if len(text) > max_chars:
        text = text[:max_chars] + " [...truncado]"
    return text


def extract_from_adjuntos(
    adjuntos: Iterable[dict],
    max_adjuntos: int = MAX_ADJUNTOS,
    max_chars_total: int = MAX_CHARS_TOTAL,
) -> str:
    """Procesa una lista de adjuntos y devuelve un bloque de texto listo
    para inyectar al user_prompt de Claude.

    Cada adjunto debe ser un dict con:
      - nombre / nombre_archivo: str
      - content / content_bytes: bytes
      - content_type: str (opcional)

    Devuelve string vacío si ninguno tiene texto extraíble.

    Formato del output:
    --- ADJUNTO 1: nombre.pdf ---
    <texto extraído>

    --- ADJUNTO 2: doc.docx ---
    <texto extraído>
    """
    if not adjuntos:
        return ""

    bloques: list[str] = []
    chars_total = 0

    for i, adj in enumerate(adjuntos):
        if i >= max_adjuntos:
            bloques.append(f"--- {len(list(adjuntos)) - max_adjuntos} adjuntos adicionales omitidos ---")
            break

        nombre = (adj.get("nombre") or adj.get("nombre_archivo") or f"adjunto_{i+1}").strip()
        content = adj.get("content") or adj.get("content_bytes") or b""
        ctype = adj.get("content_type") or ""

        if not content:
            continue

        chars_restantes = max(0, max_chars_total - chars_total)
        if chars_restantes < 200:  # no vale la pena
            break

        texto = extract_text(
            content, nombre, ctype,
            max_chars=min(MAX_CHARS_POR_ADJUNTO, chars_restantes),
        )
        if not texto:
            continue

        bloques.append(f"--- ADJUNTO {i+1}: {nombre} ---\n{texto}")
        chars_total += len(texto)

    return "\n\n".join(bloques)
